import asyncio
import logging
import os

from docx import Document
from pypdf import PdfReader

from services.translate import translate_text

logger = logging.getLogger(__name__)


async def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        return await asyncio.to_thread(_read_txt, file_path)
    if ext == ".docx":
        return await asyncio.to_thread(_read_docx, file_path)
    if ext == ".pdf":
        return await asyncio.to_thread(_read_pdf, file_path)
    raise ValueError("Qo'llab-quvvatlanmaydigan fayl turi")


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _read_docx(path: str) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _read_pdf(path: str) -> str:
    reader = PdfReader(path)
    pages_text = []
    for page in reader.pages:
        pages_text.append(page.extract_text() or "")
    return "\n".join(pages_text)


async def translate_document(file_path: str, target_lang: str, source_lang: str = "auto") -> str:
    """
    Hujjatdagi matnni o'qib, tarjima qilib, yangi .docx fayl sifatida saqlaydi.
    Qaytaradi: yaratilgan faylning yo'li.
    """
    raw_text = await extract_text(file_path)
    if not raw_text.strip():
        raise ValueError("Hujjatdan matn topilmadi (skanerlangan rasm bo'lishi mumkin).")

    translated = await translate_text(raw_text, target_lang=target_lang, source_lang=source_lang)

    out_path = os.path.splitext(file_path)[0] + f"_translated_{target_lang}.docx"

    def _write():
        doc = Document()
        doc.add_heading("Tarjima natijasi", level=1)
        for paragraph in translated.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        doc.save(out_path)

    await asyncio.to_thread(_write)
    return out_path
